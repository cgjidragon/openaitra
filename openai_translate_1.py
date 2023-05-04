import docx
import openai
import os
import re
import sys
import logging

# 设置您的OpenAI API密钥
openai.api_key = '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'

def get_docx_files(dir):
    # 初始化一个空列表，用来保存所有后缀名为docx的文件路径
    docx_files = []
    # 遍历目标文件夹中所有文件名
    for file_name in os.listdir(dir):
        # 将当前文件名组合成完整的路径
        file_path = os.path.join(dir, file_name)
        # 如果当前文件是一个目录，则递归调用本函数，继续遍历子目录
        if os.path.isdir(file_path):
            docx_files += get_docx_files(file_path)
        # 如果当前文件是后缀名为docx的文件，则将其路径添加到docx_files列表中
        elif file_name.endswith('.docx'):
            docx_files.append(file_path)
        else:
            pass
    # 返回所有后缀名为docx的文件路径
    return docx_files

def translate_docx_file(source_file_name):
    print(f'Translating file {source_file_name}...')
    target_file_name = re.sub(r'\.docx$', '_en\g<0>', source_file_name)
    source_doc = docx.Document(source_file_name)
    target_doc = docx.Document()
    for para in source_doc.paragraphs:
        translated_text = translate_text(input_text=para.text)
        target_doc.add_paragraph(translated_text)
    target_doc.save(target_file_name)
    print(f'Translated {source_file_name} to {target_file_name}')

def translate_text(input_text, source_language='zh-CN', target_language='en'):
    print(f'input_text = "{input_text[:30]}..."')
    # 使用OpenAI进行翻译
    response = openai.Completion.create(
        engine='text-davinci-003',
        prompt=f'Translate the following {source_language} text to {target_language}:\n\n{input_text}',
        max_tokens=1000,
        temperature=0.7,
        n=1,
        stop=None,
        log_level='info'
    )
    # 提取翻译结果
    translation = response.choices[0].text.strip()
    print(f'translation = "{translation[:30]}..."')
    return translation

if __name__ == '__main__':
    dir = sys.argv[1] if len(sys.argv) > 1 else os.getcwd()
    print(f'Translating files in directory {dir}...')
    for docx_file in get_docx_files(dir):
        translate_docx_file(docx_file)
    print('Translation complete!')
