from docx import Document
from docx.shared import Inches
import nltk
import re

doc = Document('example.docx')

text = []
for para in doc.paragraphs:
    text.append(para.text)
    


sentences = []
for para in text:
    sentences += re.split(r'(?<=[.?!])\s+', para.replace('\n', ' '))



max_length = 1000
tokenizer = nltk.tokenize.RegexpTokenizer(r'\w+')

for sent in sentences:
    words = tokenizer.tokenize(sent)
    if len(words) <= max_length:
        p = doc.add_paragraph()
        p.add_run(sent).bold = True
        p.add_run('\n')
    else:
        for i in range(0, len(words), max_length):
            sub_sent = ' '.join(words[i:i+max_length])
            p = doc.add_paragraph()
            p.add_run(sub_sent).bold = True
            p.add_run('\n')
