import openai

# 设置您的OpenAI API密钥
openai.api_key = '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'

import docx

# 读取Word文档的内容
def read_docx_file1(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def read_docx_file(file_path, max_tokens):
    doc = docx.Document(file_path)
    full_text = []
    current_text = ''
    for para in doc.paragraphs:
        text = para.text.strip()
        if len(current_text) + len(text) < max_tokens:
            current_text += '\n' + text if current_text else text
        else:
            full_text.append(current_text)
            current_text = text
    if current_text:
        full_text.append(current_text)
    return full_text

# 将翻译结果写入新的Word文档
def write_docx_file(file_path, text_list):
    doc = docx.Document()
    for text in text_list:
        doc.add_paragraph(text)
    doc.save(file_path)
    


# 将文本写入Word文档
def write_docx_file1(file_path, text):
    doc = docx.Document()
    doc.add_paragraph(text)
    doc.save(file_path)



def translate_text(input_text):
    # 设置您要翻译的源语言和目标语言（从中文到英文）
    source_language = 'zh-CN'
    target_language = 'en'

    # 使用OpenAI进行翻译
    response = openai.Completion.create(
        engine='text-davinci-003',
        prompt=f"Translate the following {source_language} text to {target_language}:\n\n{input_text}",
        max_tokens=1000,
        temperature=0.7,
        n=1,
        stop=None,
        #log_level='info'
    )

    # 提取翻译结果
    translation = response.choices[0].text.strip()

    return translation

# 读取txt文件内容
#with open('input.txt', 'r', encoding='utf-8') as file:
#    input_text = file.read()
#    print("input_text")
# 读取Word文档内容
input_text = read_docx_file('input.docx')
# 将文本分割成多个列表元素
max_tokens = 1000
input_text_list = read_docx_file('input.docx', max_tokens)
# 对每个列表元素进行翻译
translated_text_list = []
for text in input_text_list:
    translated_text = translate_text(text)
    translated_text_list.append(translated_text)


# 进行翻译
translated_text = translate_text(input_text)

# 打印翻译结果
print(translated_text)


# Write translated text to file
#with open('output.txt', 'w') as file:
#    file.write(translated_text)
# 将翻译结果写入新的Word文档
write_docx_file('output.docx', translated_text_list)

# Print confirmation message
print('Translation saved to output.txt')

